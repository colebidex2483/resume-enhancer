import os
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import pdfplumber
from docx import Document
import spacy
import tempfile
import uuid
from google import genai  # Using the specific import you requested
import logging
from dotenv import load_dotenv
load_dotenv()
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)
app.config['UPLOAD_FOLDER'] = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

nlp = spacy.load('en_core_web_sm')
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
gemini_client = genai.Client(api_key=GEMINI_API_KEY)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_file(file):
    if file.content_length > MAX_FILE_SIZE:
        return False, "File size exceeds 5MB limit"
    if not allowed_file(file.filename):
        return False, "Invalid file type. Only PDF, DOCX, and TXT files are allowed."
    return True, ""

def parse_file(filepath, filename):
    if filename.endswith('.pdf'):
        return parse_pdf(filepath)
    elif filename.endswith('.docx'):
        return parse_docx(filepath)
    elif filename.endswith('.txt'):
        with open(filepath, 'r') as f:
            return f.read()
    return None

def parse_pdf(filepath):
    text = ''
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ''
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
    return text

def parse_docx(filepath):
    try:
        doc = Document(filepath)
        return '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return ''

def extract_key_requirements(job_description):
    doc = nlp(job_description)
    requirements = []
    
    # Extract skills, technologies, and qualifications
    for sent in doc.sents:
        if 'experience' in sent.text.lower() or 'skill' in sent.text.lower() or 'require' in sent.text.lower():
            requirements.append(sent.text.strip())
    
    # Extract noun chunks that might represent requirements
    requirements.extend([chunk.text for chunk in doc.noun_chunks if chunk.text.lower() not in ['you', 'we', 'company']])
    
    # Extract specific entities
    requirements.extend([ent.text for ent in doc.ents if ent.label_ in ['ORG', 'PRODUCT', 'TECH']])
    
    return list(set(requirements))  # Remove duplicates

def enhance_resume_with_gemini(resume_text, job_description, key_requirements):
    prompt = f"""
    You are a professional resume writer. Your task is to optimize the following resume to perfectly match the provided job description.
    
    Job Description:
    {job_description}
    
    Key Requirements Identified:
    {', '.join(key_requirements)}
    
    Original Resume:
    {resume_text}
    
    Please rewrite the resume to:
    1. Highlight the most relevant qualifications for this specific job
    2. Incorporate keywords from the job description naturally
    3. Maintain the original structure but enhance content
    4. Keep it professional and concise
    5. Add any missing but relevant skills/experiences (mark them with [Enhanced] tag)
    
    Return ONLY the enhanced resume content, no additional commentary.
    """
    
    try:
        response = gemini_client.models.generate_content(
            model="gemini-1.5-flash",
            contents=prompt
        )
        return response.text
    except Exception as e:
        logger.error(f"Error calling Gemini API: {e}")
        return None

@app.route('/api/enhance-resume', methods=['POST'])
def enhance_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'Resume file is required'}), 400
    
    resume_file = request.files['resume']
    
    # Get job description - can be either file or text
    job_description = request.form.get('job_description_text')
    job_file = request.files.get('job_description_file')
    
    if not job_description and not job_file:
        return jsonify({'error': 'Either job description text or file is required'}), 400
    
    # Validate resume file
    is_valid, message = validate_file(resume_file)
    if not is_valid:
        return jsonify({'error': message}), 400
    
    # Process job description
    job_text = job_description  # Use text if provided
    
    # If no text but file provided, parse the file
    if not job_text and job_file:
        is_valid, message = validate_file(job_file)
        if not is_valid:
            return jsonify({'error': message}), 400
        
        # Create temp file for job description
        temp_dir = tempfile.mkdtemp()
        job_filename = secure_filename(f"job_{uuid.uuid4().hex}_{job_file.filename}")
        job_path = os.path.join(temp_dir, job_filename)
        job_file.save(job_path)
        
        try:
            job_text = parse_file(job_path, job_filename)
            if not job_text:
                return jsonify({'error': 'Failed to parse job description file'}), 400
        finally:
            # Clean up temp file
            try:
                os.unlink(job_path)
                os.rmdir(temp_dir)
            except Exception as e:
                logger.error(f"Error cleaning temp files: {e}")
    
    # Process resume file
    temp_dir = tempfile.mkdtemp()
    resume_filename = secure_filename(f"resume_{uuid.uuid4().hex}_{resume_file.filename}")
    resume_path = os.path.join(temp_dir, resume_filename)
    resume_file.save(resume_path)
    
    try:
        resume_text = parse_file(resume_path, resume_filename)
        if not resume_text:
            return jsonify({'error': 'Failed to parse resume file'}), 400
        
        # Extract key requirements from job description
        key_requirements = extract_key_requirements(job_text)
        
        # Enhance resume using Gemini
        enhanced_resume = enhance_resume_with_gemini(resume_text, job_text, key_requirements)
        
        if not enhanced_resume:
            return jsonify({'error': 'Failed to enhance resume'}), 500
        
        # Save enhanced resume as DOCX
        enhanced_filename = f"enhanced_{resume_filename.split('.')[0]}.docx"
        enhanced_path = os.path.join(temp_dir, enhanced_filename)
        
        doc = Document()
        for line in enhanced_resume.split('\n'):
            doc.add_paragraph(line.strip())
        doc.save(enhanced_path)
        
        # Return the enhanced file
        return send_file(
            enhanced_path,
            as_attachment=True,
            download_name=enhanced_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error processing files: {e}")
        return jsonify({'error': str(e)}), 500
    
    finally:
        # Clean up temp files
        try:
            for filename in os.listdir(temp_dir):
                file_path = os.path.join(temp_dir, filename)
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            os.rmdir(temp_dir)
        except Exception as e:
            logger.error(f"Error cleaning up temp files: {e}")



@app.route('/api/generate-proposal', methods=['POST'])
def generate_proposal():
    data = request.get_json()
    job_description = data.get('job_description')
    proposal_template = data.get('proposal_template')
    
    if not job_description or not proposal_template:
        return jsonify({'error': 'Both job description and proposal template are required'}), 400
    
    try:
        prompt = f"""
        You are a professional freelancer specializing in writing winning Upwork proposals. 
        Your task is to customize the following proposal template based on the provided job description.
        
        Job Description:
        {job_description}
        
        Proposal Template:
        {proposal_template}
        
        Please:
        1. Analyze the job description to understand the client's needs
        2. Customize the template by replacing placeholders with relevant content
        3. Keep the professional tone but make it personalized
        4. Highlight how the freelancer can solve the client's specific problems
        5. Keep the length appropriate (typically 3-4 paragraphs)
        
        Return ONLY the customized proposal content, no additional commentary.
        """
        
        response = gemini_client.models.generate_content(
            model="gemini-1.5-flash",
            contents=prompt
        )
        
        return jsonify({
            'proposal': response.text
        })
        
    except Exception as e:
        logger.error(f"Error generating proposal: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy'}), 200

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    port = int(os.environ.get("PORT", 5000))  # Get the port from the environment or use 5000 by default
    app.run(host='0.0.0.0', port=port)